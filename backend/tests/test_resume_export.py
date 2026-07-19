from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from uuid import UUID

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader

from app.api.routes import resumes as resumes_route
from app.main import app
from app.schemas.resume import (
    GeneratedContentProvenance,
    GeneratedResumeSection,
    ProfileEvidenceItem,
    ProfileEvidenceType,
    ProfileMatchSummary,
    RequirementMatch,
    ResumeContact,
    StructuredGeneratedResume,
    StructuredResumeRecord,
)
from app.services.export import ExportFormat, export_resume_record
from app.services.export.document_model import build_document_model
from app.services.export.export_service import EXPORT_RENDERER_VERSION
from app.services.export.filename import build_export_filename
from app.services.export.template_registry import resolve_template
from app.services.resume_store import ResumeNotFoundError, ResumeOwnershipError


USER_ID = "11111111-1111-1111-1111-111111111111"
OTHER_ID = "22222222-2222-2222-2222-222222222222"
PROFILE_ID = "33333333-3333-3333-3333-333333333333"
RESUME_ID = "44444444-4444-4444-4444-444444444444"


def evidence() -> ProfileEvidenceItem:
    return ProfileEvidenceItem(
        evidenceId="ev-infosys-api",
        evidenceType=ProfileEvidenceType.work_experience,
        sourceRecordId="experience-infosys",
        sourceLabel="Work Experience: Senior .NET Developer at Infosys",
        originalText="Built C# and ASP.NET Core REST API enhancements with SQL Server, improving release quality by 20%.",
        companyName="Infosys",
        roleTitle="Senior .NET Developer",
        strengthScore=95,
        reason="Direct evidence.",
    )


def requirement(requirement_id: str, value: str, *, matched: bool = True) -> RequirementMatch:
    return RequirementMatch(
        requirementId=requirement_id,
        requirementValue=value,
        requirementCategory="Technical",
        requirementPriority="high",
        requirementPriorityScore=90,
        classification="exact" if matched else "unmatched",
        matchScore=95 if matched else 0,
        evidence=[evidence()] if matched else [],
        isSafeToUse=matched,
    )


def profile_match() -> ProfileMatchSummary:
    return ProfileMatchSummary(
        overallMatchScore=82,
        coreRequirementScore=90,
        supportingRequirementScore=75,
        matchedRequirements=[requirement("req-c", "C#"), requirement("req-dotnet", ".NET")],
        unmatchedRequirements=[requirement("req-azure", "Azure", matched=False), requirement("req-trading", "Trading", matched=False)],
    )


def section(section_id: str, section_type: str, title: str, order: int, content, *, visible: bool = True) -> GeneratedResumeSection:
    return GeneratedResumeSection(
        sectionId=section_id,
        type=section_type,
        title=title,
        order=order,
        visible=visible,
        content=content,
        provenance=GeneratedContentProvenance(
            supportingEvidenceIds=["ev-infosys-api"] if section_type in {"summary", "experience", "skills"} else [],
            supportedRequirementIds=["req-c"] if section_type in {"summary", "experience", "skills"} else [],
            validationStatus="validated",
        ),
    )


def structured_resume(*, version: int = 2, hidden_skills: bool = False, long: bool = False) -> StructuredGeneratedResume:
    bullets = [
        {"generatedText": "Original generated bullet should not export.", "currentText": "Edited C# and ASP.NET Core REST API delivery with SQL Server, preserving the real 20% release quality metric.", "order": 2},
        {"generatedText": "Reviewed API and SQL changes with QA and architects to improve release readiness.", "order": 1},
        {"generatedText": "Deleted text should not export.", "currentText": "Deleted text should not export.", "deleted": True, "order": 3},
    ]
    if long:
        bullets.extend({"generatedText": f"Built additional C# service workflow {index} with SQL Server validation for enterprise delivery.", "order": index + 4} for index in range(80))
    return StructuredGeneratedResume(
        resumeId=RESUME_ID,
        userId=USER_ID,
        resumeName="Venu Madhav Pendurthi - Software Engineer IV",
        targetJobTitle="Software Engineer IV",
        targetCompany="Velera",
        jobDescription="Need C# .NET APIs SQL Server. Azure and Trading are not supported by profile.",
        profileId=PROFILE_ID,
        profileVersion=3,
        profileContentHash="hash",
        matchingAlgorithmVersion="match-v1",
        generationAlgorithmVersion="generation-v1",
        templateId="classic-ats",
        versionNumber=version,
        status="draft",
        matchScore=82,
        missingRequirements=["Azure", "Trading"],
        contact=ResumeContact(email="venu@example.com", phone="+12014436937", location="Hartford, CT", linkedin="https://linkedin.com/in/venu"),
        sections=[
            section("summary", "summary", "SUMMARY", 1, "Senior .NET Developer with enterprise C# API delivery experience."),
            section("skills", "skills", "TECHNICAL SKILLS", 2, [{"category": "Languages", "items": ["C#", ".NET", "SQL Server"]}], visible=not hidden_skills),
            section(
                "experience",
                "experience",
                "PROFESSIONAL EXPERIENCE",
                3,
                [{"company": "Infosys", "role": "Senior .NET Developer", "location": "Hartford, CT", "startDate": "2025-01", "endDate": "Present", "bullets": bullets}],
            ),
            section("empty-projects", "projects", "PROJECTS", 4, []),
            section("education", "education", "EDUCATION", 5, [{"degree": "MS Computer Science", "institution": "University of Central Missouri", "location": "Warrensburg, MO", "gradYear": "2022"}]),
        ],
        createdAt="2026-07-01T00:00:00+00:00",
        updatedAt="2026-07-01T00:00:00+00:00",
    )


def record(resume: StructuredGeneratedResume | None = None) -> StructuredResumeRecord:
    resume = resume or structured_resume()
    return StructuredResumeRecord(
        resumeId=resume.resume_id,
        userId=USER_ID,
        profileId=PROFILE_ID,
        profileVersion=3,
        profileContentHash="hash",
        resumeName=resume.resume_name,
        targetJobTitle=resume.target_job_title,
        targetCompany=resume.target_company,
        jobDescription=resume.job_description,
        jobAnalysisJson={},
        profileMatchJson=profile_match().model_dump(mode="json", by_alias=True),
        resumeJson=resume,
        templateId=resume.template_id,
        matchScore=resume.match_score,
        generationAlgorithmVersion=resume.generation_algorithm_version,
        status="draft",
        versionNumber=resume.version_number,
        parentResumeId="",
        createdAt=resume.created_at,
        updatedAt=resume.updated_at,
    )


def test_export_document_model_uses_structured_resume_header_visibility() -> None:
    resume = structured_resume().model_copy(
        update={
            "resume_header": {
                "fullName": "Venu Madhav Pendurthi",
                "phone": "+12014436937",
                "githubUrl": "https://github.com/venu",
            },
            "contact": ResumeContact(
                email="hidden@example.com",
                phone="+12014436937",
                location="Hidden, CT",
                linkedin="https://linkedin.com/in/hidden",
                github="https://github.com/venu",
                portfolio="https://hidden.example.com",
            ),
        }
    )

    model = build_document_model(
        resume,
        template=resolve_template("classic-ats"),
        renderer_version=EXPORT_RENDERER_VERSION,
    )

    assert model.full_name == "Venu Madhav Pendurthi"
    assert model.professional_title == ""
    assert [item.label for item in model.contact_items] == ["Phone", "GitHub"]
    assert [item.value for item in model.contact_items] == ["+12014436937", "https://github.com/venu"]


def test_export_document_model_preserves_rendered_skill_groups_for_pdf_and_docx() -> None:
    resume = structured_resume()
    skills_section = next(item for item in resume.sections if item.type == "skills")
    skills_section.content = [
        {
            "category": "Programming Languages",
            "items": ["C#", "SQL/T-SQL"],
            "sourceSkillIds": ["skill-csharp", "skill-sql"],
            "renderingPolicyVersion": "skills-rendering-v1",
        },
        {
            "category": "Backend Frameworks & Tools",
            "items": ["ASP.NET Core", "REST APIs"],
            "sourceSkillIds": ["skill-aspnet-core", "skill-rest-apis"],
            "renderingPolicyVersion": "skills-rendering-v1",
        },
    ]

    template = resolve_template("classic-ats")
    pdf_model = build_document_model(resume, template=template, renderer_version=EXPORT_RENDERER_VERSION)
    docx_model = build_document_model(resume, template=template, renderer_version=EXPORT_RENDERER_VERSION)

    assert pdf_model.sections[1].content == docx_model.sections[1].content
    assert [(group.category, group.items) for group in pdf_model.sections[1].content] == [
        ("Programming Languages", ["C#", "SQL/T-SQL"]),
        ("Backend Frameworks & Tools", ["ASP.NET Core", "REST APIs"]),
    ]


def pdf_text(content: bytes) -> str:
    return "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)


def docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def test_document_model_uses_current_text_and_falls_back_to_generated_text() -> None:
    model = build_document_model(record().resume_json, template=resolve_template("classic-ats"), renderer_version=EXPORT_RENDERER_VERSION)
    bullets = model.sections[2].content[0].bullets

    assert bullets[0].text.startswith("Reviewed API")
    assert "Edited C#" in bullets[1].text
    assert "Original generated" not in " ".join(bullet.text for bullet in bullets)


def test_document_model_omits_hidden_empty_deleted_and_internal_ids() -> None:
    model = build_document_model(structured_resume(hidden_skills=True), template=resolve_template("classic-ats"), renderer_version=EXPORT_RENDERER_VERSION)
    payload = model.model_dump(mode="json")

    assert "TECHNICAL SKILLS" not in [item.title for item in model.sections]
    assert "PROJECTS" not in [item.title for item in model.sections]
    assert "Deleted text" not in str(payload)
    assert "ev-infosys-api" not in str(payload)


def test_document_model_is_deterministic_and_does_not_mutate_source_json() -> None:
    resume = structured_resume()
    before = resume.model_dump(mode="json")
    template = resolve_template("classic-ats")
    first = build_document_model(resume, template=template, renderer_version=EXPORT_RENDERER_VERSION).model_dump(exclude={"export_metadata"})
    second = build_document_model(resume, template=template, renderer_version=EXPORT_RENDERER_VERSION).model_dump(exclude={"export_metadata"})

    assert first == second
    assert resume.model_dump(mode="json") == before


def test_build_export_filename_sanitizes_unsafe_values() -> None:
    filename = build_export_filename(full_name="Venu / Madhav", target_role="Senior .NET: Developer", company="Velera<>", version_number=2, extension="pdf")

    assert filename == "Venu_Madhav_Senior_NET_Developer_Velera_v2.pdf"


def test_export_valid_pdf_contains_extractable_resume_text() -> None:
    result = export_resume_record(record(), export_format=ExportFormat.pdf)
    text = pdf_text(result.content)

    assert result.content.startswith(b"%PDF")
    assert result.content_type == "application/pdf"
    assert "VENU MADHAV PENDURTHI" in text
    assert "C#" in text
    assert ".NET" in text
    assert "20%" in text
    assert "40%" not in text
    assert "Azure" not in text
    assert "Trading" not in text


def test_export_allows_stale_evidence_references_as_warnings() -> None:
    resume = structured_resume()
    stale_section = resume.sections[0].model_copy(
        update={
            "provenance": resume.sections[0].provenance.model_copy(
                update={"supporting_evidence_ids": ["stale-evidence-id"], "supported_requirement_ids": ["req-c"]}
            )
        }
    )
    resume = resume.model_copy(update={"sections": [stale_section, *resume.sections[1:]]})

    result = export_resume_record(record(resume), export_format=ExportFormat.pdf)

    assert result.content.startswith(b"%PDF")
    assert any("unknown evidence" in warning.lower() for warning in result.warnings)


def test_export_allows_duplicate_bullets_as_warnings() -> None:
    resume = structured_resume()
    experience_section = resume.sections[2]
    experience = deepcopy(experience_section.content[0])
    duplicate_text = "Reviewed API and SQL changes with QA and architects to improve release readiness."
    experience["bullets"] = [
        {"generatedText": duplicate_text},
        {"generatedText": duplicate_text},
    ]
    resume = resume.model_copy(
        update={
            "sections": [
                *resume.sections[:2],
                experience_section.model_copy(update={"content": [experience]}),
                *resume.sections[3:],
            ]
        }
    )

    result = export_resume_record(record(resume), export_format=ExportFormat.docx)

    assert result.content[:2] == b"PK"
    assert any("duplicate generated bullet" in warning.lower() for warning in result.warnings)


def test_pdf_preserves_section_and_bullet_order_and_hidden_sections() -> None:
    result = export_resume_record(record(structured_resume(hidden_skills=True)), export_format=ExportFormat.pdf)
    text = pdf_text(result.content)

    assert text.index("SUMMARY") < text.index("PROFESSIONAL EXPERIENCE") < text.index("EDUCATION")
    assert text.index("Reviewed API") < text.index("Edited C#")
    assert "TECHNICAL SKILLS" not in text


def test_multi_page_pdf_is_valid_and_not_truncated() -> None:
    result = export_resume_record(record(structured_resume(long=True)), export_format=ExportFormat.pdf)
    reader = PdfReader(BytesIO(result.content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    assert len(reader.pages) >= 2
    assert "service workflow 79" in text


def test_export_valid_docx_is_editable_ooxml_with_styles() -> None:
    result = export_resume_record(record(), export_format=ExportFormat.docx)
    document = Document(BytesIO(result.content))
    text = docx_text(result.content)
    styles = {style.name for style in document.styles}

    assert result.content_type.endswith("wordprocessingml.document")
    assert result.content[:2] == b"PK"
    assert "VENU MADHAV PENDURTHI" in text
    assert "Edited C# and ASP.NET Core" in text
    assert "20%" in text
    assert "Azure" not in text
    assert "Trading" not in text
    assert "Resume Bullet" in styles
    assert any(paragraph.style.name == "Resume Bullet" for paragraph in document.paragraphs)


def test_docx_preserves_order_and_hides_sections() -> None:
    text = docx_text(export_resume_record(record(structured_resume(hidden_skills=True)), export_format=ExportFormat.docx).content)

    assert text.index("SUMMARY") < text.index("PROFESSIONAL EXPERIENCE") < text.index("EDUCATION")
    assert text.index("Reviewed API") < text.index("Edited C#")
    assert "TECHNICAL SKILLS" not in text


def test_historical_version_exports_historical_content() -> None:
    historical = structured_resume(version=1).model_copy(update={"resume_name": "Venu Madhav Pendurthi - Historical"})
    historical.sections[0].content = "Historical summary content."
    result = export_resume_record(record(historical), export_format=ExportFormat.pdf)

    assert "Historical summary content" in pdf_text(result.content)
    assert result.resume_version == 1


def test_invalid_template_id_returns_validation_error() -> None:
    with pytest.raises(ValueError):
        export_resume_record(record(), export_format=ExportFormat.pdf, template_id="javascript:bad")


def test_api_pdf_export_headers_and_body(monkeypatch) -> None:
    async def fake_session():
        yield None

    async def fake_get_resume(session, user_id, resume_id):
        assert user_id == UUID(USER_ID)
        return record()

    app.dependency_overrides[resumes_route.db_session] = fake_session
    app.dependency_overrides[resumes_route.optional_current_user_id] = lambda: UUID(USER_ID)
    monkeypatch.setattr(resumes_route, "get_resume", fake_get_resume)
    try:
        response = TestClient(app).get(f"/api/resumes/{RESUME_ID}/export/pdf")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.headers["cache-control"] == "private, no-store"
    assert response.headers["x-export-renderer-version"] == EXPORT_RENDERER_VERSION
    assert response.content.startswith(b"%PDF")


def test_api_docx_export_uses_selected_version(monkeypatch) -> None:
    async def fake_session():
        yield None

    async def fake_get_resume(session, user_id, resume_id):
        return record(structured_resume(version=3))

    app.dependency_overrides[resumes_route.db_session] = fake_session
    app.dependency_overrides[resumes_route.optional_current_user_id] = lambda: UUID(USER_ID)
    monkeypatch.setattr(resumes_route, "get_resume", fake_get_resume)
    try:
        response = TestClient(app).get(f"/api/resumes/{RESUME_ID}/export/docx")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    assert response.headers["x-resume-version"] == "3"
    assert response.content[:2] == b"PK"


def test_api_export_rejects_other_user(monkeypatch) -> None:
    async def fake_session():
        yield None

    async def fake_get_resume(session, user_id, resume_id):
        raise ResumeOwnershipError("You do not have access to this resume.")

    app.dependency_overrides[resumes_route.db_session] = fake_session
    app.dependency_overrides[resumes_route.optional_current_user_id] = lambda: UUID(OTHER_ID)
    monkeypatch.setattr(resumes_route, "get_resume", fake_get_resume)
    try:
        response = TestClient(app).get(f"/api/resumes/{RESUME_ID}/export/pdf")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 403


def test_api_export_invalid_resume_id_returns_404(monkeypatch) -> None:
    async def fake_session():
        yield None

    async def fake_get_resume(session, user_id, resume_id):
        raise ResumeNotFoundError("Resume not found.")

    app.dependency_overrides[resumes_route.db_session] = fake_session
    app.dependency_overrides[resumes_route.optional_current_user_id] = lambda: UUID(USER_ID)
    monkeypatch.setattr(resumes_route, "get_resume", fake_get_resume)
    try:
        response = TestClient(app).get("/api/resumes/not-a-real-id/export/pdf")
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 404


def test_api_export_sanitizes_requested_filename(monkeypatch) -> None:
    async def fake_session():
        yield None

    async def fake_get_resume(session, user_id, resume_id):
        return record()

    app.dependency_overrides[resumes_route.db_session] = fake_session
    app.dependency_overrides[resumes_route.optional_current_user_id] = lambda: UUID(USER_ID)
    monkeypatch.setattr(resumes_route, "get_resume", fake_get_resume)
    try:
        response = TestClient(app).get(f"/api/resumes/{RESUME_ID}/export/pdf?filename=..%2Funsafe%3Cname%3E")
    finally:
        app.dependency_overrides.clear()

    assert 'filename="unsafe_name.pdf"' in response.headers["content-disposition"]
