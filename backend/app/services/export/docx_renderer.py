from __future__ import annotations

from io import BytesIO

from app.services.export.document_model import (
    CertificationItem,
    EducationItem,
    ExperienceItem,
    ProjectItem,
    ResumeDocumentModel,
    SkillGroup,
)
from app.services.export.template_registry import ResumeExportTemplate


def render_docx(model: ResumeDocumentModel, template: ResumeExportTemplate, *, paper_size: str = "letter") -> bytes:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    if paper_size.lower() == "a4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    top, bottom, left, right = template.margins
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

    configure_styles(document, template)
    document.core_properties.title = model.full_name
    document.core_properties.subject = "Resume"
    document.core_properties.author = model.full_name
    document.core_properties.keywords = f"resume;{model.template_id};{model.export_metadata.renderer_version}"

    name = document.add_paragraph(style="Resume Name")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.add_run(model.full_name.upper())
    if model.professional_title:
        title = document.add_paragraph(style="Resume Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(model.professional_title)
    if model.contact_items:
        contact = document.add_paragraph(style="Resume Contact")
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(" | ".join(item.value for item in model.contact_items))

    for export_section in model.sections:
        add_section_heading(document, export_section.title, template)
        if export_section.type == "summary":
            document.add_paragraph(str(export_section.content), style="Resume Body")
        elif export_section.type == "skills":
            for item in export_section.content:
                group = item if isinstance(item, SkillGroup) else SkillGroup.model_validate(item)
                paragraph = document.add_paragraph(style="Resume Body")
                label = paragraph.add_run(f"{group.category}: ")
                label.bold = True
                paragraph.add_run(", ".join(group.items))
        elif export_section.type == "experience":
            for item in export_section.content:
                exp = item if isinstance(item, ExperienceItem) else ExperienceItem.model_validate(item)
                add_role_block(document, exp.company, exp.role, exp.location, exp.start_date, exp.end_date, [bullet.text for bullet in exp.bullets])
        elif export_section.type == "projects":
            for item in export_section.content:
                project = item if isinstance(item, ProjectItem) else ProjectItem.model_validate(item)
                paragraph = document.add_paragraph(style="Resume Role Heading")
                paragraph.add_run(" | ".join(part for part in [project.name, project.org, project.link] if part)).bold = True
                if project.technologies:
                    tech = document.add_paragraph(style="Resume Body")
                    label = tech.add_run("Technologies: ")
                    label.bold = True
                    tech.add_run(", ".join(project.technologies))
                for bullet in project.bullets:
                    document.add_paragraph(bullet.text, style="Resume Bullet")
        elif export_section.type == "education":
            for item in export_section.content:
                education = item if isinstance(item, EducationItem) else EducationItem.model_validate(item)
                document.add_paragraph(" | ".join(part for part in [education.degree, education.institution, education.location, education.grad_year, education.gpa] if part), style="Resume Body")
        elif export_section.type == "certifications":
            for item in export_section.content:
                certification = item if isinstance(item, CertificationItem) else CertificationItem.model_validate(item)
                document.add_paragraph(" | ".join(part for part in [certification.name, certification.issuer, certification.issued_date, certification.expiry_date] if part), style="Resume Body")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def configure_styles(document, template: ResumeExportTemplate) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    font_name = "Helvetica" if template.font_name == "Helvetica" else "Times New Roman"
    accent = RGBColor.from_string(template.accent_color.lstrip("#"))
    styles = document.styles
    styles["Normal"].font.name = font_name
    styles["Normal"].font.size = Pt(template.body_size)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def style(name: str, size: int, *, bold: bool = False, italic: bool = False, color=None, space_after: float = 2):
        paragraph_style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        paragraph_style.font.name = font_name
        paragraph_style.font.size = Pt(size)
        paragraph_style.font.bold = bold
        paragraph_style.font.italic = italic
        if color:
            paragraph_style.font.color.rgb = color
        paragraph_style.paragraph_format.space_after = Pt(space_after)
        paragraph_style.paragraph_format.space_before = Pt(0)
        paragraph_style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        return paragraph_style

    style("Resume Name", template.name_size, bold=True, space_after=1)
    style("Resume Title", template.body_size + 1, space_after=1)
    style("Resume Contact", max(template.body_size - 1, 8), space_after=5)
    style("Resume Section Heading", template.section_size, bold=True, color=accent, space_after=2)
    style("Resume Role Heading", template.body_size, bold=True, space_after=0)
    style("Resume Body", template.body_size, space_after=2)
    bullet = style("Resume Bullet", template.body_size, space_after=1.5)
    if "List Bullet" in styles:
        bullet.base_style = styles["List Bullet"]
    bullet.paragraph_format.left_indent = Inches(0.24)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)


def add_section_heading(document, title: str, template: ResumeExportTemplate) -> None:
    paragraph = document.add_paragraph(style="Resume Section Heading")
    paragraph.add_run(title.upper())
    add_bottom_border(paragraph, template.accent_color)


def add_role_block(document, company: str, role: str, location: str, start: str, end: str, bullets: list[str]) -> None:
    period = " - ".join(part for part in [start, end] if part)
    heading = document.add_paragraph(style="Resume Role Heading")
    heading.add_run(" | ".join(part for part in [company, period] if part)).bold = True
    subheading = document.add_paragraph(style="Resume Body")
    subheading.add_run(" | ".join(part for part in [role, location] if part)).italic = True
    for bullet in bullets:
        document.add_paragraph(bullet, style="Resume Bullet")


def add_bottom_border(paragraph, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.lstrip("#"))
    p_bdr.append(bottom)
