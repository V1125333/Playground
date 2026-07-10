from html import escape
from io import BytesIO

from app.schemas.resume import ResumeContent


def resume_filename(resume: ResumeContent, extension: str) -> str:
    name = "_".join(part for part in resume.name.split() if part) or "resume"
    role = "_".join(part for part in resume.title.split()[:4] if part) or "generated"
    safe = "".join(char if char.isalnum() or char in {"_", "-"} else "_" for char in f"{name}_{role}")
    return f"{safe}.{extension}"


def build_docx(resume: ResumeContent) -> bytes:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.60)
    section.right_margin = Inches(0.60)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    if "Resume Bullet" not in styles:
        bullet_style = styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.base_style = styles["Normal"]
    else:
        bullet_style = styles["Resume Bullet"]
    bullet_style.paragraph_format.left_indent = Inches(0.25)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.15)
    bullet_style.paragraph_format.space_after = Pt(3.5)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(resume.name.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)

    if resume.title:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(resume.title)
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(12)

    contacts = contact_items(resume)
    if contacts:
        contact = document.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_contact_runs(contact, contacts)

    add_section(document, "SUMMARY")
    document.add_paragraph(resume.summary or "")

    if resume.skills:
        add_section(document, "TECHNICAL SKILLS")
        for group in resume.skills:
            paragraph = document.add_paragraph()
            label = paragraph.add_run(f"{group.category}: ")
            label.bold = True
            paragraph.add_run(", ".join(group.items))

    if resume.experience:
        add_section(document, "PROFESSIONAL EXPERIENCE")
        for item in resume.experience:
            add_two_column_line(document, item.company, period(item.start_date, item.end_date), bold_left=True)
            add_two_column_line(document, item.role, item.location, bold_left=True, italic_right=True, space_after=2)
            for bullet in item.bullets:
                paragraph = document.add_paragraph(style="Resume Bullet")
                paragraph.add_run("\u2022\t")
                paragraph.add_run(bullet)

    if resume.projects:
        add_section(document, "PROJECTS")
        for project in resume.projects:
            add_two_column_line(document, " | ".join(filter(None, [project.name, project.org])), project.link, bold_left=True)
            for bullet in project.bullets:
                paragraph = document.add_paragraph(style="Resume Bullet")
                paragraph.add_run("\u2022\t")
                paragraph.add_run(bullet)
            if project.technologies:
                paragraph = document.add_paragraph()
                label = paragraph.add_run("Technologies: ")
                label.bold = True
                paragraph.add_run(", ".join(project.technologies))

    if resume.education:
        add_section(document, "EDUCATION")
        for education in resume.education:
            add_two_column_line(document, education.degree, education.grad_year, bold_left=True)
            paragraph = document.add_paragraph(" | ".join(filter(None, [education.institution, education.location, education.gpa])))
            if education.location:
                paragraph.runs[-1].italic = True

    if resume.certifications:
        add_section(document, "CERTIFICATIONS")
        for certification in resume.certifications:
            document.add_paragraph(certification_line(certification))

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf(resume: ResumeContent) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer

    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        leftMargin=0.60 * inch,
        rightMargin=0.60 * inch,
    )
    base = getSampleStyleSheet()
    normal = ParagraphStyle("ResumeNormal", parent=base["Normal"], fontName="Times-Roman", fontSize=11, leading=13)
    bold = ParagraphStyle("ResumeBold", parent=normal, fontName="Times-Bold")
    name_style = ParagraphStyle("ResumeName", parent=bold, alignment=TA_CENTER, fontSize=22, leading=24)
    centered = ParagraphStyle("ResumeCentered", parent=normal, alignment=TA_CENTER, fontSize=11, leading=13)
    heading = ParagraphStyle("ResumeHeading", parent=bold, fontSize=11, leading=12, spaceBefore=6, spaceAfter=4)
    company = ParagraphStyle("ResumeCompany", parent=normal, fontName="Times-Bold", fontSize=11, leading=13, spaceAfter=0)
    role = ParagraphStyle("ResumeRole", parent=normal, fontName="Times-Bold", fontSize=11, leading=13, spaceAfter=1)
    bullet = ParagraphStyle("ResumeBullet", parent=normal, leftIndent=0, firstLineIndent=0, spaceBefore=0, spaceAfter=0)
    bullet_mark = ParagraphStyle("ResumeBulletMark", parent=normal, alignment=TA_RIGHT, leftIndent=0, firstLineIndent=0, spaceAfter=0)
    right = ParagraphStyle("ResumeRight", parent=normal, alignment=TA_RIGHT)

    story = [
        Paragraph(escape(resume.name.upper()), name_style),
        Paragraph(escape(resume.title), centered) if resume.title else Spacer(0, 0),
    ]
    contacts = contact_items(resume)
    if contacts:
        story.append(Paragraph(escape("  |  ".join(contacts)), centered))

    def section(title: str) -> None:
        story.append(Spacer(0, 8))
        story.append(Paragraph(title, heading))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black, spaceAfter=4))

    section("SUMMARY")
    story.append(Paragraph(escape(resume.summary or ""), normal))

    if resume.skills:
        section("TECHNICAL SKILLS")
        for group in resume.skills:
            story.append(Paragraph(f"<b>{escape(group.category)}:</b> {escape(', '.join(group.items))}", normal))

    if resume.experience:
        section("PROFESSIONAL EXPERIENCE")
        for item in resume.experience:
            block = [
                two_column_paragraph(item.company, period(item.start_date, item.end_date), company, right),
                two_column_paragraph(item.role, item.location, role, right, italic_right=True),
            ]
            block.extend(bullet_row(bullet_text, bullet, bullet_mark) for bullet_text in item.bullets)
            story.append(KeepTogether(block))

    if resume.projects:
        section("PROJECTS")
        for project in resume.projects:
            block = [two_column_paragraph(" | ".join(filter(None, [project.name, project.org])), project.link, company, right)]
            block.extend(bullet_row(bullet_text, bullet, bullet_mark) for bullet_text in project.bullets)
            if project.technologies:
                block.append(Paragraph(f"<b>Technologies:</b> {escape(', '.join(project.technologies))}", normal))
            story.append(KeepTogether(block))

    if resume.education:
        section("EDUCATION")
        for education in resume.education:
            story.append(two_column_paragraph(education.degree, education.grad_year, normal, right))
            story.append(Paragraph(escape(" | ".join(filter(None, [education.institution, education.location, education.gpa]))), normal))

    if resume.certifications:
        section("CERTIFICATIONS")
        for certification in resume.certifications:
            story.append(Paragraph(escape(certification_line(certification)), normal))

    document.build(story)
    return output.getvalue()


def contact_items(resume: ResumeContent) -> list[str]:
    return [item for item in [
        resume.contact.phone,
        resume.contact.email,
        resume.contact.location,
        resume.contact.linkedin,
        resume.contact.portfolio,
    ] if item]


def certification_line(certification) -> str:
    details = [
        certification.issuer,
        f"Issued {certification.issued_date}" if certification.issued_date else "",
        f"Expires {certification.expiry_date}" if certification.expiry_date else "",
    ]
    suffix = " | ".join(item for item in details if item)
    return f"{certification.name} - {suffix}" if suffix else certification.name


def period(start: str, end: str) -> str:
    return " - ".join(item for item in [start, end] if item)


def add_section(document, title: str) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    add_bottom_border(paragraph)


def add_bottom_border(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "111111")
    p_bdr.append(bottom)


def add_two_column_line(document, left: str, right: str, *, bold_left: bool = False, italic_right: bool = False, space_after: int = 0) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.space_after = Pt(space_after)
    left_run = paragraph.add_run(left)
    left_run.bold = bold_left
    if right:
        paragraph.add_run("\t")
        right_run = paragraph.add_run(right)
        right_run.italic = italic_right


def add_contact_runs(paragraph, contacts: list[str]) -> None:
    for index, item in enumerate(contacts):
        if index:
            paragraph.add_run("  |  ")
        paragraph.add_run(item)


def two_column_paragraph(left: str, right_text: str, normal, right_style, italic_right: bool = False):
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Table, TableStyle

    right_markup = f"<i>{escape(right_text)}</i>" if italic_right and right_text else escape(right_text)
    return Table(
        [[Paragraph(escape(left), normal), Paragraph(right_markup, right_style)]],
        colWidths=[5.65 * inch, 1.65 * inch],
        style=TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("LINEBELOW", (0, 0), (-1, -1), 0, colors.white),
        ]),
    )


def bullet_row(text: str, bullet_style, marker_style):
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, Table, TableStyle

    return Table(
        [[Paragraph("\u2022", marker_style), Paragraph(escape(text), bullet_style)]],
        colWidths=[0.24 * inch, 7.06 * inch],
        style=TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (0, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
            ("LINEBELOW", (0, 0), (-1, -1), 0, colors.white),
        ]),
    )


from reportlab.platypus import Paragraph
